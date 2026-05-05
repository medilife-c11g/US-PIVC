#!/usr/bin/env python3
"""
Round 2: apply M3 / M5 / M6 / M7 fixes to POMIN_v2 in-place.

User-confirmed decisions:
  M3: phone ext = 71297 (correct cover letter; fix title page)
  M5: Refosco NOS = 5/9 (per authoritative risk_of_bias.md; reflect catheter-length confounding)
  M6: Jen-Tao Wang affiliation = Department of Surgery, Show Chwan Memorial Hospital (original)
  M7: RCT-only = RR 1.03 (95% CI 0.66–1.60), p = 0.91 (random-effects DL — internally consistent w/ primary)

For M6 we restructure affiliations (insert Surgery as #3; push Research Assistant Center to #4,
Nursing to #5; Po-Ming Chen now uses 4, 5).
"""

from pathlib import Path
from copy import deepcopy
from docx import Document
from docx.oxml.ns import qn

POMIN_V2 = Path("/Users/chencc/Research/US-PIVC/POMIN_v2_2026-05-05")


def replace_in_paragraph_xml(para, old, new):
    el = para._element
    t_elements = el.findall(".//" + qn("w:t"))
    if not t_elements:
        return False
    full = "".join((t.text or "") for t in t_elements)
    if old not in full:
        return False
    new_full = full.replace(old, new)
    t_elements[0].text = new_full
    for t in t_elements[1:]:
        t.text = ""
    return True


def fix_manuscript_round2():
    p = POMIN_V2 / "manuscript 20260505.docx"
    doc = Document(p)
    edits = []

    # --- M3: phone extension on title page ---
    for para in doc.paragraphs[:25]:
        if replace_in_paragraph_xml(para, "ext:71299", "ext:71297"):
            edits.append("M3 phone ext fixed")
            break

    # --- M6: Jen-Tao Wang affiliation 1 → 3 (Surgery, Show Chwan) ---
    # Need to update author line "Jen-Tao Wang1" → "Jen-Tao Wang3"
    # AND update affiliations list (insert Surgery as #3; renumber)
    for para in doc.paragraphs[:25]:
        # Author byline currently: "...Jen-Tao Wang1, Po-Ming Chen3, 4, Chia-Ching Chen1*"
        # After: "...Jen-Tao Wang3, Po-Ming Chen4, 5, Chia-Ching Chen1*"
        if "Jen-Tao Wang" in para.text and "Po-Ming Chen" in para.text:
            ok1 = replace_in_paragraph_xml(para, "Jen-Tao Wang1", "Jen-Tao Wang3")
            # Po-Ming Chen 3, 4 → 4, 5 (after inserting Surgery as new #3)
            ok2 = (
                replace_in_paragraph_xml(para, "Po-Ming Chen3, 4", "Po-Ming Chen4, 5")
                or replace_in_paragraph_xml(para, "Po-Ming Chen 3, 4", "Po-Ming Chen 4, 5")
                or replace_in_paragraph_xml(para, "Po-Ming Chen3,4", "Po-Ming Chen4,5")
            )
            if ok1 and ok2:
                edits.append("M6 author-line affiliations renumbered")
            break

    # Update affiliation list paragraphs: rewrite #3 and add #4, #5 if needed
    # Original POMIN had:
    #   1 EM Chang Bing Show Chwan
    #   2 Gastroenterology Chang Bing Show Chwan
    #   3 Research Assistant Center, Show Chwan Memorial Hospital
    #   4 Department of Nursing, Central Taiwan University of Science and Technology, Taichung
    # Target:
    #   1 EM Chang Bing Show Chwan
    #   2 Gastroenterology Chang Bing Show Chwan
    #   3 Department of Surgery, Show Chwan Memorial Hospital  ← NEW (was #3 RAC)
    #   4 Research Assistant Center, Show Chwan Memorial Hospital  ← shifted from #3
    #   5 Department of Nursing, Central Taiwan University of Science and Technology  ← shifted from #4
    for para in doc.paragraphs[:25]:
        # POMIN aff #3 paragraph contains "3 Research Assistant Center"
        if "3" in para.text and "Research Assistant Center" in para.text:
            replace_in_paragraph_xml(
                para,
                "3 Research Assistant Center, Show Chwan Memorial Hospital, Changhua, Taiwan",
                "3 Department of Surgery, Show Chwan Memorial Hospital, Changhua, Taiwan",
            )
            # also handle without trailing ", Changhua, Taiwan"
            replace_in_paragraph_xml(
                para,
                "3 Research Assistant Center, Show Chwan Memorial Hospital",
                "3 Department of Surgery, Show Chwan Memorial Hospital",
            )
            edits.append("M6 affiliation #3 → Surgery")
        if "4 Department of Nursing" in para.text:
            replace_in_paragraph_xml(
                para,
                "4 Department of Nursing, Central Taiwan University of Science and Technology, Taichung, Taiwan",
                "4 Research Assistant Center, Show Chwan Memorial Hospital, Changhua, Taiwan; "
                "5 Department of Nursing, Central Taiwan University of Science and Technology, Taichung, Taiwan",
            )
            replace_in_paragraph_xml(
                para,
                "4 Department of Nursing, Central Taiwan University of Science and Technology",
                "4 Research Assistant Center, Show Chwan Memorial Hospital, Changhua, Taiwan; "
                "5 Department of Nursing, Central Taiwan University of Science and Technology, Taichung, Taiwan",
            )
            edits.append("M6 affiliation #4 split → #4 RAC + #5 Nursing")

    # --- M5: Refosco NOS row (Table 2 Part B) ---
    # Find row containing 'Refosco' across all tables and replace cell text 6/9 → 5/9,
    # and "Catheter type confounding" → "Catheter length confounding (64 mm vs 19–32 mm)".
    # Also fix Selection cells from 3 → 2 (per risk_of_bias.md: 2 selection stars due to major
    # catheter-length selection bias).
    refosco_edits = 0
    for table in doc.tables:
        for row in table.rows:
            if "Refosco" not in row.cells[0].text:
                continue
            for cell in row.cells:
                # Walk all <w:t> in this cell, do replacements
                for para in cell.paragraphs:
                    if replace_in_paragraph_xml(para, "6/9", "5/9"):
                        refosco_edits += 1
                    if replace_in_paragraph_xml(para, "Catheter type confounding",
                                                "Catheter length confounding (64 mm vs 19–32 mm)"):
                        refosco_edits += 1
                    # Selection: change 3 → 2 ONLY if cell contains exactly "3" (Selection score)
                    # Be careful — many cells have "3" — only target Selection stars.
                    # We'll handle this via dedicated cell-level approach below.
            # Now selectively change Selection stars: cells before the Comparability cell.
            # Heuristic: in Part B header row (row 2 of table), Selection cells are 1-2 (or 2-3
            # depending on Study cell merge count). The "Refosco row" duplicates: cell[1]=cell[2]='3'
            # (Selection), cell[3]=cell[4]='1' (Comparability), cell[5]=cell[6]='2' (Outcome),
            # cell[7]=cell[8]='5/9' (Total — already updated), then 9-11 = Key concern.
            # We need cell[1] and cell[2] = '3' → '2'.
            for ci in (1, 2):
                if ci < len(row.cells):
                    cell = row.cells[ci]
                    for para in cell.paragraphs:
                        if replace_in_paragraph_xml(para, "3", "2"):
                            refosco_edits += 1
            # Also fix Total cell: '5/9' is now wrong (actually total: 2+1+2 = 5, so 5/9 is right).
            # No further change needed for Total.
            break  # only one Refosco row
    if refosco_edits:
        edits.append(f"M5 Refosco NOS: {refosco_edits} cell text edits (3 stars → 2; 6/9 → 5/9; key concern wording)")

    # --- M7: RCT-only p = 0.88 → p = 0.91 (random-effects DL for consistency) ---
    # Two paragraphs reference this (paragraph 99 and 126 from earlier grep)
    m7_count = 0
    for para in doc.paragraphs:
        # Variant 1: "p = 0.88" in catheter-failure context
        if replace_in_paragraph_xml(para, "RR 1.03, 95% CI 0.66–1.61; p = 0.88",
                                    "RR 1.03, 95% CI 0.66–1.60; p = 0.91 (random-effects DerSimonian-Laird, consistent with primary model)"):
            m7_count += 1
        # Variant: just "p = 0.88" alone
        elif "p = 0.88" in para.text and "RCT" in para.text and "1.03" in para.text:
            replace_in_paragraph_xml(para, "p = 0.88",
                                     "p = 0.91 (random-effects DerSimonian-Laird, consistent with primary model)")
            replace_in_paragraph_xml(para, "1.03, 95% CI 0.66–1.61", "1.03, 95% CI 0.66–1.60")
            m7_count += 1
        # Variant in paragraph 126 — narrative mention without p-value, but CI 0.66-1.61 still wrong
        elif ("RCT-only estimate was flat" in para.text or "RCT-only estimate was essentially flat" in para.text) and "0.66–1.61" in para.text:
            replace_in_paragraph_xml(para, "0.66–1.61", "0.66–1.60")
            m7_count += 1
    if m7_count:
        edits.append(f"M7: RCT-only p adjusted to 0.91 (random-effects DL); CI 1.61→1.60 in {m7_count} para(s)")

    doc.save(p)
    return edits


if __name__ == "__main__":
    print(f"Working in: {POMIN_V2}\n")
    edits = fix_manuscript_round2()
    for e in edits:
        print(f"  ✓ {e}")
    print(f"\nTotal: {len(edits)} edits")
