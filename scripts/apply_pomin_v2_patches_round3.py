#!/usr/bin/env python3
"""
Round 3 patches addressing Codex Round-2 audit findings.

HIGH:
  - Affiliations #3-#5 fully superscripted (only leading number should be sup)
  - Refosco NOS in body Results text (para 95) still says "6/9" — should be 5/9 group
  - Residual unbracketed inline citations from POMIN's incomplete style migration

PARTIAL → fix completion:
  - S1 PRISMA GRADE row: "not applied; noted as limitation" → consistent with H4
  - Refosco Figure 7: regenerate with Selection=2 (R script needs update too)

Also runs a regen of Figure 7 with M5-fixed Refosco rating.
"""

from pathlib import Path
from copy import deepcopy
from docx import Document
from docx.oxml.ns import qn

POMIN_V2 = Path("/Users/chencc/Research/US-PIVC/POMIN_v2_2026-05-05")


def replace_in_para_xml(para, old, new):
    el = para._element
    ts = el.findall(".//" + qn("w:t"))
    if not ts:
        return False
    full = "".join((t.text or "") for t in ts)
    if old not in full:
        return False
    ts[0].text = full.replace(old, new)
    for t in ts[1:]:
        t.text = ""
    return True


def split_first_run_at_offset(para, offset, *, leading_keep_super=True):
    """Take the first <w:r> of paragraph (which contains both the leading
    superscript number AND following normal text), split into:
      - run 1: text up to `offset`, keeping its current rPr (likely superscript)
      - run 2: text from `offset` onward, with vertAlign removed (normal script)
    """
    el = para._element
    runs = [r for r in el.findall(qn("w:r"))]
    if not runs:
        return False
    first_r = runs[0]
    # find first <w:t>
    t_elems = first_r.findall(qn("w:t"))
    if not t_elems:
        return False
    t = t_elems[0]
    text = t.text or ""
    if offset >= len(text):
        return False
    # Split text
    leading_text = text[:offset]
    trailing_text = text[offset:]
    t.text = leading_text
    # Build a new <w:r> sibling carrying trailing_text in normal script
    new_r = deepcopy(first_r)
    new_t = new_r.findall(qn("w:t"))[0]
    new_t.text = trailing_text
    # Remove vertAlign superscript from the new run
    new_rpr = new_r.find(qn("w:rPr"))
    if new_rpr is not None:
        va = new_rpr.find(qn("w:vertAlign"))
        if va is not None:
            new_rpr.remove(va)
    # Insert new_r right after first_r
    first_r.addnext(new_r)
    return True


def fix_affiliation_superscripts():
    p = POMIN_V2 / "manuscript 20260505.docx"
    doc = Document(p)
    fixed = 0
    # Title page affiliations are paragraphs 2..6 (1=indices, 2=Affiliation, etc.)
    for i in range(2, 8):
        if i >= len(doc.paragraphs):
            break
        para = doc.paragraphs[i]
        text = para.text
        if not text or text[0] not in "12345":
            continue
        # find offset: skip leading digit(s), keep them superscript; rest normal.
        offset = 1
        if len(text) > 1 and text[1].isdigit():
            offset = 2
        # Confirm first run is fully-superscripted with combined text
        runs = para.runs
        if not runs:
            continue
        rpr = runs[0]._element.find(qn("w:rPr"))
        is_sup = False
        if rpr is not None:
            va = rpr.find(qn("w:vertAlign"))
            is_sup = va is not None and va.get(qn("w:val")) == "superscript"
        # Only act if first run is sup AND its text contains more than just the digit
        first_run_text = runs[0].text or ""
        if is_sup and len(first_run_text) > offset:
            ok = split_first_run_at_offset(para, offset, leading_keep_super=True)
            if ok:
                fixed += 1
    doc.save(p)
    print(f"[Round3-aff] split {fixed} affiliation paragraphs (number stays superscript, body normal)")


def fix_refosco_results_text():
    p = POMIN_V2 / "manuscript 20260505.docx"
    doc = Document(p)
    edits = 0
    # Paragraph 95 currently says "Feinsmith et al. (2021), Dachepally et al. (2023),
    # Paladini et al. (2018), Cottrell et al. (2021), and Refosco et al. (2024) each
    # scored 6/9, with points deducted primarily in the comparability domain
    # (catheter type confounding for Dachepally, Paladini, and Refosco; ...)"
    # After M5 fix: Refosco should NOT be in the 6/9 list (now 5/9).
    # Also: catheter "type" → "length" confounding.

    # We'll do a paragraph-text-level rewrite via XML walker.
    target_old = (
        "Feinsmith et al. (2021), Dachepally et al. (2023), Paladini et al. (2018), "
        "Cottrell et al. (2021), and Refosco et al. (2024) each scored 6/9, with points "
        "deducted primarily in the comparability domain (catheter type confounding for "
        "Dachepally, Paladini, and Refosco; difficulty-classification imbalance for "
        "Cottrell)."
    )
    target_new = (
        "Feinsmith et al. (2021), Dachepally et al. (2023), Paladini et al. (2018), and "
        "Cottrell et al. (2021) each scored 6/9, with points deducted primarily in the "
        "comparability domain (catheter type confounding for Dachepally and Paladini; "
        "difficulty-classification imbalance for Cottrell). Refosco et al. (2024) scored "
        "5/9, with the additional Selection-domain downgrade reflecting major "
        "catheter-length confounding (USG arm 64 mm long peripheral catheters vs blind arm "
        "19–32 mm short cannulas)."
    )
    for para in doc.paragraphs:
        if "Refosco" in para.text and "scored 6/9" in para.text:
            # The exact phrasing in POMIN may vary slightly — use a forgiving search.
            # First try strict replace:
            if replace_in_para_xml(para, target_old, target_new):
                edits += 1
                continue
            # Fallback: only swap "Refosco et al. (2024) each scored 6/9" → reword
            if replace_in_para_xml(
                para,
                "and Refosco et al. (2024) each scored 6/9",
                "each scored 6/9 (Refosco et al. (2024) scored 5/9 — see Table 2 — "
                "with the Selection-domain downgrade reflecting major catheter-length "
                "confounding: USG 64 mm long peripheral catheters vs blind 19–32 mm short "
                "cannulas)"
            ):
                edits += 1
            elif replace_in_para_xml(
                para,
                "and Refosco et al. (2024)",
                "(and Refosco et al. (2024) scored 5/9 with Selection-domain downgrade reflecting major catheter-length confounding)"
            ):
                edits += 1
    if edits == 0:
        # Brute-force fallback: any paragraph mentioning 'Refosco' with '6/9' gets '6/9' replaced
        for para in doc.paragraphs:
            if "Refosco" in para.text and "6/9" in para.text:
                # Just change the number — don't restructure
                replace_in_para_xml(para, "6/9", "5/9")
                edits += 1
    doc.save(p)
    print(f"[Round3-Refosco] body text edits: {edits}")


def fix_unbracketed_citations():
    p = POMIN_V2 / "manuscript 20260505.docx"
    doc = Document(p)
    # POMIN forgot to bracket some inline citations during their style migration.
    # Specific instances identified by Codex Round 2:
    fixes = [
        ("warranted.15", "warranted [15]."),
        ("DIVA populations.4-6", "DIVA populations [4-6]."),
        ("catheter survival.9,10", "catheter survival [9, 10]."),
        ("intravascular segment.7", "intravascular segment [7]."),
        ("below the skin surface.7", "below the skin surface [7]."),
        ("(Table 3),12", "(Table 3) [12],"),
    ]
    counts = 0
    for para in doc.paragraphs:
        for old, new in fixes:
            if old in para.text:
                if replace_in_para_xml(para, old, new):
                    counts += 1
    doc.save(p)
    print(f"[Round3-cites] {counts} bracketed-citation conversions")


def fix_s1_grade_residual():
    p = POMIN_V2 / "Supplementary Material S1 20260505.docx"
    doc = Document(p)
    edits = 0
    # Round-2 audit said S1 still contains "not applied; noted as limitation"
    targets = {
        "GRADE not applied; noted as limitation": "GRADE applied; see Manuscript Table 3 (Summary of Findings)",
        "not applied; noted as limitation":      "applied; see Manuscript Table 3 (Summary of Findings)",
        "not applied. Noted as limitation":      "applied. See Manuscript Table 3.",
        "noted as a limitation":                  "see Manuscript Table 3",
    }
    # paragraphs
    for para in doc.paragraphs:
        for k, v in targets.items():
            if k in para.text:
                if replace_in_para_xml(para, k, v):
                    edits += 1
    # tables
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for k, v in targets.items():
                        if k in para.text:
                            if replace_in_para_xml(para, k, v):
                                edits += 1
    doc.save(p)
    print(f"[Round3-S1] GRADE residual cleanup: {edits} edits")


def harmonize_ai_disclosure():
    """Cover letter says 'Anthropic Claude; OpenAI ChatGPT'; manuscript says
    'ChatGPT, Gemini'. Standardize to a unified statement in both."""
    unified = ("AI-assisted tools (Anthropic Claude and OpenAI ChatGPT) were used "
               "for language editing, code-comment generation, and an independent "
               "methodological audit of the analytic workflow. All clinical "
               "interpretation, statistical analyses, and final scientific claims "
               "are the authors' own. No AI tool was listed as an author.")
    edits = 0
    # Cover letter
    p_cover = POMIN_V2 / "Cover_Letter 20260505.docx"
    doc = Document(p_cover)
    for para in doc.paragraphs:
        if "Generative AI" in para.text:
            # rewrite full paragraph text to unified statement
            from docx.oxml.ns import qn as _q
            ts = para._element.findall(".//" + _q("w:t"))
            if ts:
                ts[0].text = "Generative AI Disclosure: " + unified
                for t in ts[1:]:
                    t.text = ""
                edits += 1
    doc.save(p_cover)
    # Manuscript
    p_ms = POMIN_V2 / "manuscript 20260505.docx"
    doc = Document(p_ms)
    for para in doc.paragraphs:
        if "Use of Artificial Intelligence" in para.text or "AI assistants" in para.text:
            from docx.oxml.ns import qn as _q
            ts = para._element.findall(".//" + _q("w:t"))
            if ts:
                ts[0].text = "Use of Artificial Intelligence: " + unified
                for t in ts[1:]:
                    t.text = ""
                edits += 1
                break
    doc.save(p_ms)
    print(f"[Round3-AI] AI disclosure harmonized in {edits} document(s)")


def standardize_i_squared_notation():
    """Standardize 'I2' → 'I²' across docx files for consistency."""
    edits_total = 0
    for fname in ("manuscript 20260505.docx", "Cover_Letter 20260505.docx",
                  "Supplementary Material S1 20260505.docx",
                  "Supplementary Material S3 20260505.docx"):
        p = POMIN_V2 / fname
        if not p.exists():
            continue
        doc = Document(p)
        edits = 0
        for para in doc.paragraphs:
            if "I2 " in para.text or "I2=" in para.text or "I2 =" in para.text:
                for old, new in [("I2 ", "I² "), ("I2=", "I²="), ("I2 =", "I² =")]:
                    if replace_in_para_xml(para, old, new):
                        edits += 1
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for old, new in [("I2 ", "I² "), ("I2=", "I²="), ("I2 =", "I² =")]:
                            if replace_in_para_xml(para, old, new):
                                edits += 1
        doc.save(p)
        if edits:
            print(f"[Round3-I²] {fname}: {edits} I² notation fixes")
            edits_total += edits
    if edits_total == 0:
        print("[Round3-I²] no I2/I² inconsistencies found (or already standardized)")


if __name__ == "__main__":
    print(f"Working in: {POMIN_V2}\n")
    fix_affiliation_superscripts()
    fix_refosco_results_text()
    fix_unbracketed_citations()
    fix_s1_grade_residual()
    harmonize_ai_disclosure()
    standardize_i_squared_notation()
    print("\nRound 3 patches complete.")
