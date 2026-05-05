#!/usr/bin/env python3
"""
Apply Codex audit H1-H5 fixes to POMIN package, producing POMIN_v2_2026-05-05/.

Edits use python-docx; preserves Word document structure where possible.
"""

import shutil
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from copy import deepcopy

POMIN_V2 = Path("/Users/chencc/Research/US-PIVC/POMIN_v2_2026-05-05")

# ======================================================================
# H5 — Cover letter: fix CI string and add AI disclosure (M1 also)
# ======================================================================
def fix_cover_letter():
    p = POMIN_V2 / "Cover_Letter 20260505.docx"
    doc = Document(p)
    edits = 0
    for para in doc.paragraphs:
        for run in para.runs:
            if "1.00–1.51" in run.text:
                run.text = run.text.replace(
                    "1.00–1.51",
                    "0.99–1.51 (precise: 0.995–1.508; lower bound rounds to 1.00 at 2 dp but the unrounded interval crosses unity)",
                )
                edits += 1
            if "1.00-1.51" in run.text:  # ASCII hyphen variant
                run.text = run.text.replace(
                    "1.00-1.51",
                    "0.99-1.51 (precise: 0.995-1.508; lower bound rounds to 1.00 at 2 dp but the unrounded interval crosses unity)",
                )
                edits += 1
    # M1: append AI disclosure paragraph before signature block
    # find the "Sincerely," paragraph and insert before it
    sincerely_idx = None
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip().startswith("Sincerely"):
            sincerely_idx = i
            break
    if sincerely_idx is not None:
        # insert a new paragraph before Sincerely
        new_para_text = (
            "Generative AI Disclosure: AI-assisted tools (Anthropic Claude; OpenAI ChatGPT) "
            "were used for language editing, code-comment generation, and an independent "
            "methodological audit of the analytic workflow. All clinical interpretation, "
            "statistical analyses, and final scientific claims are the authors' own. No AI "
            "tool was listed as an author. Full disclosure appears in the manuscript Methods."
        )
        sincerely_para = doc.paragraphs[sincerely_idx]
        # insert before the sincerely paragraph by manipulating XML
        new_p = sincerely_para._element.addprevious(deepcopy(sincerely_para._element))
        # the duplicated paragraph is now BEFORE sincerely; clear its runs and set text
        new_para = doc.paragraphs[sincerely_idx]  # the duplicated one
        for r in new_para.runs:
            r.text = ""
        new_para.runs[0].text = new_para_text if new_para.runs else None
        if not new_para.runs:
            new_para.add_run(new_para_text)
        else:
            new_para.runs[0].text = new_para_text
        edits += 1
    doc.save(p)
    print(f"[H5+M1] Cover letter: {edits} edits applied")


# ======================================================================
# H2 + H3 + M2 + M3 + M6 — Title page (within manuscript.docx)
# ======================================================================
def replace_in_paragraph(para, old, new):
    """Replace text inside a paragraph that may span multiple runs. Concatenates
    runs into the first run and clears the rest. Format of first run preserved."""
    full = "".join(r.text for r in para.runs)
    if old not in full:
        return False
    new_full = full.replace(old, new)
    if not para.runs:
        return False
    para.runs[0].text = new_full
    for r in para.runs[1:]:
        r.text = ""
    return True


def collapse_paragraph_text(para, new_text):
    """Force-set paragraph text by writing all into first run and clearing others."""
    if not para.runs:
        para.add_run(new_text)
        return True
    para.runs[0].text = new_text
    for r in para.runs[1:]:
        r.text = ""
    return True


def replace_anywhere_in_paragraph_xml(para, old, new):
    """Replace text inside paragraph by walking ALL <w:t> elements (including
    those inside <w:hyperlink>). Combines text across w:t elements first to
    catch matches that span elements. Edits in place."""
    from docx.oxml.ns import qn
    el = para._element
    t_elements = el.findall(".//" + qn("w:t"))
    if not t_elements:
        return False
    full = "".join((t.text or "") for t in t_elements)
    if old not in full:
        return False
    new_full = full.replace(old, new)
    # write everything into the first w:t, clear the rest
    t_elements[0].text = new_full
    for t in t_elements[1:]:
        t.text = ""
    return True


def fix_manuscript_title_page_and_contributions():
    p = POMIN_V2 / "manuscript 20260505.docx"
    doc = Document(p)
    edits = 0
    for para in doc.paragraphs:
        # H2: affiliation #5 doesn't exist; PMC was "4, 5" — fix to "3, 4"
        # Also paragraph contains affiliation list — we need to add aff #5 OR fix PMC to #3,4
        # Approach: replace "Po-Ming Chen4, 5" / "Po-Ming Chen 4, 5" -> "Po-Ming Chen3, 4"
        for variant in [
            "Po-Ming Chen4, 5", "Po-Ming Chen4,5", "Po-Ming Chen 4, 5",
            "Po-Ming Chen 4 , 5", "Po-Ming Chen 4,5", "Po-Ming Chen 4 ,5",
        ]:
            if replace_anywhere_in_paragraph_xml(para, variant, "Po-Ming Chen3, 4"):
                edits += 1
                break
        # M2: email typo medlife -> medilife on title page (handles hyperlinks)
        if replace_anywhere_in_paragraph_xml(para, "medlife.c11g@gmail.com", "medilife.c11g@gmail.com"):
            edits += 1
    # H3: Authors' Contributions section — POMIN dropped this entirely.
    # Insert ICMJE-compliant Authors' Contributions + Data Availability AFTER
    # "Conflicts of interest:" and BEFORE "Use of Artificial Intelligence:".
    coi_idx = None
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip().startswith("Conflicts of interest"):
            coi_idx = i
            break
    if coi_idx is not None:
        coi_para = doc.paragraphs[coi_idx]
        # Insert two new paragraphs AFTER COI by walking XML siblings
        contrib_text = (
            "Authors' Contributions: T.A.L. and Y.L.W. contributed equally and share co-first "
            "authorship. C.C.C. (corresponding author) conceived and designed the study, "
            "performed the literature search, data extraction, statistical analysis, and "
            "drafted the manuscript. T.A.L. independently performed title/abstract screening, "
            "data extraction, and risk of bias assessment, and critically reviewed the "
            "manuscript. Y.L.W. independently performed title/abstract screening and data "
            "extraction, and critically reviewed the manuscript. J.T.W. assisted with data "
            "verification and critically reviewed the manuscript. P.M.C. (research project "
            "assistant) supported data verification, manuscript formatting, supplementary "
            "file preparation, and submission package assembly, and critically reviewed the "
            "manuscript. All authors read and approved the final version."
        )
        data_avail_text = (
            "Availability of Data and Materials: The datasets generated and analysed during "
            "the current study, including the final 14-study extraction CSV, the R analysis "
            "script, and all derived forest-plot / GRADE / risk-of-bias outputs, are "
            "available from the corresponding author on reasonable request."
        )
        from copy import deepcopy
        # Duplicate the COI paragraph element twice as templates for new paragraphs
        for new_text in (contrib_text, data_avail_text):
            new_p = deepcopy(coi_para._element)
            coi_para._element.addnext(new_p)
            # The newly inserted element now sits immediately after coi_para;
            # for next iteration we want the NEXT one to come AFTER this new one,
            # so update coi_para reference
            from docx.text.paragraph import Paragraph
            new_para_obj = Paragraph(new_p, coi_para._parent)
            replace_anywhere_in_paragraph_xml(new_para_obj, "Conflicts of interest: None declared.", new_text)
            # If text was COI text, replace; if other variant, just collapse all w:t into new_text
            current_text = new_para_obj.text
            if current_text != new_text:
                # force-set
                from docx.oxml.ns import qn
                t_elements = new_para_obj._element.findall(".//" + qn("w:t"))
                if t_elements:
                    t_elements[0].text = new_text
                    for t in t_elements[1:]:
                        t.text = ""
            coi_para = new_para_obj  # so next insert chains correctly
            edits += 1
    doc.save(p)
    print(f"[H2+H3+M2] Manuscript title page/contributions: {edits} edits applied")


# ======================================================================
# H4 — Add GRADE Table 3 to manuscript
# ======================================================================
def add_grade_table_to_manuscript():
    p = POMIN_V2 / "manuscript 20260505.docx"
    doc = Document(p)
    grade_rows = [
        ("Outcome", "Studies (k)", "Participants", "Effect Estimate", "Certainty (GRADE)", "Reasons for Downgrade"),
        ("Catheter failure", "4 (2 RCTs, 2 cohorts)", "3,404",
         "RR 1.23 (95% CI 0.99–1.51); p=0.056",
         "⊕⊕◯◯ Low",
         "Risk of bias (open-label, non-blinded outcome assessment); imprecision (CI crosses null)"),
        ("Dwell time", "3 (2 RCTs, 1 cohort)", "929",
         "Not pooled (I²=91.9%)",
         "⊕◯◯◯ Very low",
         "Inconsistency (extreme heterogeneity); risk of bias; imprecision"),
        ("Infiltration", "2 (1 RCT, 1 cohort)", "2,778",
         "RR 0.68 (95% CI 0.12–3.83); NS",
         "⊕◯◯◯ Very low",
         "Risk of bias; inconsistency; severe imprecision (very wide CI)"),
        ("Extravasation", "3 (2 RCTs, 1 cohort)", "29,682",
         "Not pooled (I²=95.7%)",
         "⊕◯◯◯ Very low",
         "Inconsistency (extreme heterogeneity); risk of bias; indirectness"),
        ("Time-to-event-adjacent (failure / survival)", "2 cohorts", "43,800",
         "Not pooled — Feinsmith adj HR 0.91 (0.87–0.95); Shokoohi 72-h RR 1.26 (0.88–1.80)",
         "⊕◯◯◯ Very low",
         "Different effect measures (HR vs RR); risk of bias; conflicting directions; imprecision"),
    ]
    # Find the paragraph containing "Table 2." (the Risk of Bias table heading) — we will add Table 3 after Table 2
    target_idx = None
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip().startswith("Part B: Newcastle-Ottawa Scale"):
            target_idx = i
            # find end of Table 2: walk forward until next paragraph whose text starts with "Figure" or "Table" or empty
            break
    # Easiest: just append to end of doc as Table 3
    doc.add_paragraph()
    heading = doc.add_paragraph()
    heading_run = heading.add_run(
        "Table 3. GRADE Summary of Findings — Post-insertion outcomes of USG vs landmark PIVCs"
    )
    heading_run.bold = True
    table = doc.add_table(rows=1, cols=6)
    try:
        table.style = "Table Grid"
    except KeyError:
        pass  # use default style
    hdr_cells = table.rows[0].cells
    for j, h in enumerate(grade_rows[0]):
        hdr_cells[j].text = h
        for run in hdr_cells[j].paragraphs[0].runs:
            run.bold = True
    for row in grade_rows[1:]:
        cells = table.add_row().cells
        for j, v in enumerate(row):
            cells[j].text = v
    footnote = doc.add_paragraph()
    footnote_run = footnote.add_run(
        "GRADE certainty levels: ⊕⊕⊕⊕ High; ⊕⊕⊕◯ Moderate; ⊕⊕◯◯ Low; ⊕◯◯◯ Very low. "
        "All outcomes were downgraded for risk of bias because included studies are predominantly "
        "open-label with non-blinded outcome assessment, which is unavoidable for this comparison "
        "(ultrasound vs landmark cannot be blinded to the operator). Pooled estimates with I²>75% "
        "were not synthesized per pre-specified protocol; narrative synthesis is provided in Results. "
        "Abbreviations: CI, confidence interval; HR, hazard ratio; NS, not statistically significant; "
        "PIVC, peripheral intravenous catheter; RCT, randomized controlled trial; RR, risk ratio; "
        "USG, ultrasound-guided."
    )
    for run in footnote.runs:
        run.italic = True
        run.font.size = Pt(9)
    doc.save(p)
    print(f"[H4] Added GRADE Table 3 (5 outcome rows + footnote)")


# ======================================================================
# H1 + S1 contradiction — Supplementary S3 + S1
# ======================================================================
def fix_supplementary_s3():
    """
    POMIN S3 lists 6 excluded studies; manuscript Methods says 5. Authoritative
    reconciliation (per /Users/chencc/Research/US-PIVC/output/screening_consensus.md
    line 93–96): the two Kleidon overlap reports are a single conceptual exclusion
    of the EPIC trial (already counted as the included Kleidon 2025 primary).
    Also the ICU-nurses unpublished report was excluded at second-pass T/A
    adjudication, not at full-text. Restore S3 to 5 entries to match Methods.
    """
    p = POMIN_V2 / "Supplementary Material S3 20260505.docx"
    doc = Document(p)
    edits = 0
    # Update header text "n = 6" -> "n = 5"
    for para in doc.paragraphs:
        for run in para.runs:
            if "(n = 6)" in run.text:
                run.text = run.text.replace("(n = 6)", "(n = 5)")
                edits += 1
            if "(n=6)" in run.text:
                run.text = run.text.replace("(n=6)", "(n=5)")
                edits += 1
    # Also check tables for "n = 6"
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if "(n = 6)" in run.text:
                            run.text = run.text.replace("(n = 6)", "(n = 5)")
                            edits += 1
                        if "(n=6)" in run.text:
                            run.text = run.text.replace("(n=6)", "(n=5)")
                            edits += 1
    # Remove ICU-nurses row from the excluded studies table (the row whose Detail
    # contains "ICU Nurses" or whose Reason contains "Manuscript not published")
    # Find tables and identify candidate row
    rows_removed = 0
    for table in doc.tables:
        rows_to_remove = []
        for row in table.rows:
            row_text = " | ".join(cell.text for cell in row.cells)
            # Identify the ICU-nurses unpublished row specifically
            if "ICU Nurses" in row_text and "Not published" in row_text:
                rows_to_remove.append(row)
        for r in rows_to_remove:
            r._element.getparent().remove(r._element)
            rows_removed += 1
    # Also renumber remaining excluded-rows so the column "#" reads 1..5 not 1..6
    # We'll do this in the first table that contains numbered exclusion rows.
    for table in doc.tables:
        # detect: first column is "#" header
        if not table.rows:
            continue
        if table.rows[0].cells[0].text.strip() == "#":
            counter = 0
            for row in table.rows[1:]:
                first_cell = row.cells[0].text.strip()
                if first_cell.isdigit():
                    counter += 1
                    # Replace the first cell text to renumber
                    # Find first paragraph of first cell, edit its text
                    cell = row.cells[0]
                    # Clear existing paragraphs runs, set to counter
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.text = ""
                    if cell.paragraphs:
                        if cell.paragraphs[0].runs:
                            cell.paragraphs[0].runs[0].text = str(counter)
                        else:
                            cell.paragraphs[0].add_run(str(counter))
                    edits += 1
            break  # only first numbered table
    doc.save(p)
    print(f"[H1] S3: removed {rows_removed} ICU-nurses row, applied {edits} text/renumber edits")


def fix_supplementary_s1():
    """
    H4 sub-issue: S1 PRISMA checklist may state GRADE was 'not formally assessed',
    contradicting manuscript. Update S1 to confirm GRADE WAS performed and is in Table 3.
    """
    p = POMIN_V2 / "Supplementary Material S1 20260505.docx"
    doc = Document(p)
    edits = 0
    targets = {
        "not formally assessed": "performed; see Manuscript Table 3 (GRADE Summary of Findings)",
        "GRADE not assessed": "GRADE assessed; see Manuscript Table 3",
        "Not formally assessed": "Performed; see Manuscript Table 3 (GRADE Summary of Findings)",
    }
    for para in doc.paragraphs:
        for run in para.runs:
            for k, v in targets.items():
                if k in run.text:
                    run.text = run.text.replace(k, v)
                    edits += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        for k, v in targets.items():
                            if k in run.text:
                                run.text = run.text.replace(k, v)
                                edits += 1
    doc.save(p)
    print(f"[H4-S1] Supplementary S1: {edits} GRADE statement edits")


# ======================================================================
# Run all
# ======================================================================
if __name__ == "__main__":
    print(f"Working in: {POMIN_V2}")
    fix_cover_letter()
    fix_manuscript_title_page_and_contributions()
    add_grade_table_to_manuscript()
    fix_supplementary_s3()
    fix_supplementary_s1()
    print("\nDone. Verify each .docx in Word.")
